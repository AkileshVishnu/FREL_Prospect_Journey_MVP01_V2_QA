"""
generate_demo_doc.py
--------------------
Generates the FIPSAR Intelligence — Demo Q&A Documentation (.docx)
Run: python generate_demo_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────
NAVY        = RGBColor(0x1B, 0x3A, 0x6B)   # headings / table headers
TEAL        = RGBColor(0x00, 0x7A, 0x7A)   # section labels
GOLD        = RGBColor(0xD4, 0xA0, 0x17)   # accent / question number
DARK_GRAY   = RGBColor(0x2D, 0x2D, 0x2D)   # body
MID_GRAY    = RGBColor(0x60, 0x60, 0x60)   # sub-labels
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE  = RGBColor(0xE8, 0xF0, 0xFB)   # table alt row
LIGHT_GREEN = RGBColor(0xE6, 0xF4, 0xEA)   # insight box bg
CODE_BG     = RGBColor(0xF4, 0xF4, 0xF4)   # SQL block bg


# ─────────────────────────────────────────────
# Low-level XML helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="AAAAAA", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, enabled in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        if enabled:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), sz)
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
        else:
            el.set(qn('w:val'), 'none')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def para_border_bottom(para, color="CCCCCC", sz="6"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─────────────────────────────────────────────
# Convenience writers
# ─────────────────────────────────────────────

def add_run(para, text, bold=False, italic=False, color=None, size=None, font="Calibri"):
    run = para.add_run(text)
    run.font.name = font
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def heading1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(4)
    add_run(para, text, bold=True, color=NAVY, size=17, font="Calibri Light")
    para_border_bottom(para, color="1B3A6B", sz="8")
    return para


def heading2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, text, bold=True, color=TEAL, size=12)
    return para


def label_para(doc, label, value, label_color=MID_GRAY, value_color=DARK_GRAY, size=10.5):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    add_run(para, label, bold=True, color=label_color, size=size)
    add_run(para, value, bold=False, color=value_color, size=size)
    return para


def body_para(doc, text, size=10.5, color=DARK_GRAY, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(3)
    add_run(para, text, color=color, size=size)
    return para


def bullet_para(doc, text, size=10.5):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    add_run(para, text, color=DARK_GRAY, size=size)
    return para


def sql_block(doc, sql_text):
    """Render a SQL code block in a 1-cell table with monospaced font."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "F4F4F4")
    set_cell_borders(cell, color="BBBBBB")
    cell.width = Inches(6.5)
    para = cell.paragraphs[0]
    para.paragraph_format.left_indent = Inches(0.1)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(sql_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    doc.add_paragraph()   # spacing after block


def expected_table(doc, headers, rows):
    """Render a styled results preview table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1B3A6B")
        set_cell_borders(cell, color="1B3A6B")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        run.font.name = "Calibri"

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "EAF2FF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, color="CCCCCC", sz="2")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()


def insight_box(doc, text):
    """Teal-left-bordered insight callout using a borderless table trick."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Left accent cell
    accent = tbl.cell(0, 0)
    set_cell_bg(accent, "007A7A")
    accent.width = Inches(0.08)
    # Content cell
    content = tbl.cell(0, 1)
    set_cell_bg(content, "E6F4EA")
    set_cell_borders(content, color="007A7A", sz="4")
    p = content.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("💡 " + text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    doc.add_paragraph()


def divider(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para_border_bottom(para, color="DDDDDD", sz="4")


# ─────────────────────────────────────────────
# Document data — 10 Q&A entries
# ─────────────────────────────────────────────

QUESTIONS = [

    # ── Q1 ──────────────────────────────────────────────────────────────────
    {
        "num": 1,
        "question": "Give me a full pipeline funnel summary — how many leads came in and how many became valid prospects?",
        "category": "Funnel Overview",
        "business_insight": (
            "Of the 800 raw leads ingested in Q1 2026, 758 became validated Gold-layer prospects "
            "(94.75% end-to-end conversion). Two rejection stages shrink the pool: 37 records failed "
            "the PHI mastering quality gate (null emails, null phones, no-consent violations) and 5 "
            "duplicate record IDs were suppressed in the Silver layer. "
            "The 94.75% conversion rate is strong and indicates a healthy intake pipeline."
        ),
        "expected_headers": ["Stage", "Table", "Record Count", "Drop-Off"],
        "expected_rows": [
            ["Raw Intake",   "STG_PROSPECT_INTAKE",  "800",  "—"],
            ["PHI Mastered", "PHI_PROSPECT_MASTER",  "763",  "−37 (DQ failures)"],
            ["Silver",       "SLV_PROSPECT_MASTER",  "758",  "−5 (duplicates)"],
            ["Gold Dim",     "DIM_PROSPECT",         "758",  "—"],
            ["Gold Fact",    "FACT_PROSPECT_INTAKE",  "758",  "—"],
        ],
        "validation_sql": """\
-- Full pipeline funnel count validation
SELECT 'STG_PROSPECT_INTAKE'  AS stage, COUNT(*) AS cnt FROM QA_FIPSAR_PHI_HUB.STAGING.STG_PROSPECT_INTAKE
UNION ALL
SELECT 'PHI_PROSPECT_MASTER', COUNT(*) FROM QA_FIPSAR_PHI_HUB.PHI_CORE.PHI_PROSPECT_MASTER
UNION ALL
SELECT 'SLV_PROSPECT_MASTER', COUNT(*) FROM QA_FIPSAR_DW.SILVER.SLV_PROSPECT_MASTER
UNION ALL
SELECT 'DIM_PROSPECT',        COUNT(*) FROM QA_FIPSAR_DW.GOLD.DIM_PROSPECT
UNION ALL
SELECT 'FACT_PROSPECT_INTAKE',COUNT(*) FROM QA_FIPSAR_DW.GOLD.FACT_PROSPECT_INTAKE
ORDER BY cnt DESC;""",
        "insight_callout": (
            "5 records flagged as DUPLICATE_RECORD_ID in Silver were originally seeded as test fixtures "
            "(FIP0000000001–005). They are correctly suppressed and their rejection is logged in "
            "QA_FIPSAR_AUDIT.PIPELINE_AUDIT.DQ_REJECTION_LOG."
        ),
    },

    # ── Q2 ──────────────────────────────────────────────────────────────────
    {
        "num": 2,
        "question": "What are the most common rejection reasons and how many records were affected at each stage?",
        "category": "Data Quality & Rejections",
        "business_insight": (
            "The pipeline rejected 42 records total across two stages. NULL_EMAIL (31 records) is the "
            "dominant failure at the PHI gate — indicating that a significant portion of form submissions "
            "or data imports are missing email addresses. This is an actionable issue: strengthening "
            "form validation at the point of capture would directly improve prospect yield."
        ),
        "expected_headers": ["Rejection Stage", "Reason", "Count"],
        "expected_rows": [
            ["PHI_LOAD",    "NULL_EMAIL",          "31"],
            ["PHI_LOAD",    "NULL_PHONE_NUMBER",   "4"],
            ["PHI_LOAD",    "NO_CONSENT",          "2"],
            ["SILVER_LOAD", "DUPLICATE_RECORD_ID", "5"],
            ["TOTAL",       "—",                   "42"],
        ],
        "validation_sql": """\
-- Rejection breakdown by stage and reason
SELECT
    REJECTION_STAGE,
    REJECTION_REASON,
    COUNT(*) AS rejected_count
FROM QA_FIPSAR_AUDIT.PIPELINE_AUDIT.DQ_REJECTION_LOG
GROUP BY REJECTION_STAGE, REJECTION_REASON
ORDER BY rejected_count DESC;""",
        "insight_callout": (
            "The 2 NO_CONSENT rejections (Jackson Stewart, Michael Campbell) represent a HIPAA/privacy "
            "compliance gate. These records were correctly excluded from all downstream tables — "
            "verifiable by confirming FIP0000000037 and FIP0000000064 are absent from PHI_PROSPECT_MASTER."
        ),
    },

    # ── Q3 ──────────────────────────────────────────────────────────────────
    {
        "num": 3,
        "question": "Which marketing channels are driving the most prospect volume, and which have the lowest rejection rate?",
        "category": "Channel Performance",
        "business_insight": (
            "Website and Facebook are the top two channels by raw lead volume. However, channel quality "
            "matters as much as volume — channels with higher NULL_EMAIL rates (often social scrapes) "
            "have lower net-to-gross ratios. Referral consistently delivers cleaner records because "
            "contact details are provided directly by the referrer."
        ),
        "expected_headers": ["Channel", "Raw Leads", "Valid Prospects", "Net-Gross %"],
        "expected_rows": [
            ["Website",      "~175", "~167", "~95%"],
            ["Facebook",     "~165", "~154", "~93%"],
            ["Instagram",    "~140", "~134", "~96%"],
            ["Referral",     "~120", "~117", "~98%"],
            ["Campaign App", "~115", "~107", "~93%"],
            ["Survey",       "~85",  "~79",  "~93%"],
        ],
        "validation_sql": """\
-- Channel-level lead vs prospect volume
SELECT
    s.CHANNEL,
    COUNT(DISTINCT s.RECORD_ID)                               AS raw_leads,
    COUNT(DISTINCT f.RECORD_ID)                               AS valid_prospects,
    ROUND(COUNT(DISTINCT f.RECORD_ID) * 100.0
          / NULLIF(COUNT(DISTINCT s.RECORD_ID), 0), 1)        AS net_gross_pct
FROM QA_FIPSAR_PHI_HUB.STAGING.STG_PROSPECT_INTAKE     s
LEFT JOIN QA_FIPSAR_DW.GOLD.FACT_PROSPECT_INTAKE       f
       ON s.RECORD_ID = f.RECORD_ID
GROUP BY s.CHANNEL
ORDER BY raw_leads DESC;""",
        "insight_callout": (
            "Referral channel shows the highest net-gross conversion (~98%) because contact details "
            "are actively provided — making it the highest-quality acquisition source despite lower volume."
        ),
    },

    # ── Q4 ──────────────────────────────────────────────────────────────────
    {
        "num": 4,
        "question": "What is the overall SFMC email engagement — open rate, click-to-open rate, and unsubscribe rate?",
        "category": "Email Engagement KPIs",
        "business_insight": (
            "Email engagement is very strong. With a 75.7% open rate against 4,286 sends and an 88.9% "
            "click-to-open rate, the campaign content is resonating well with the prospect audience. "
            "The unsubscribe rate should be monitored to ensure journey fatigue is not building up "
            "over the multi-stage nurture sequence."
        ),
        "expected_headers": ["Metric", "Value", "Formula"],
        "expected_rows": [
            ["Total Sent",          "4,286",  "COUNT(SENT events)"],
            ["Total Opened",        "3,246",  "COUNT(OPEN events)"],
            ["Total Clicked",       "2,885",  "COUNT(CLICK events)"],
            ["Open Rate",           "75.7%",  "OPEN / SENT"],
            ["Click-to-Open Rate",  "88.9%",  "CLICK / OPEN"],
            ["Click Rate",          "67.3%",  "CLICK / SENT"],
        ],
        "validation_sql": """\
-- Overall SFMC email KPIs
SELECT
    COUNT_IF(EVENT_TYPE = 'SENT')   AS total_sent,
    COUNT_IF(EVENT_TYPE = 'OPEN')   AS total_opened,
    COUNT_IF(EVENT_TYPE = 'CLICK')  AS total_clicked,
    ROUND(COUNT_IF(EVENT_TYPE='OPEN')  * 100.0
          / NULLIF(COUNT_IF(EVENT_TYPE='SENT'),0), 1)  AS open_rate_pct,
    ROUND(COUNT_IF(EVENT_TYPE='CLICK') * 100.0
          / NULLIF(COUNT_IF(EVENT_TYPE='OPEN'),0), 1)  AS click_to_open_pct,
    ROUND(COUNT_IF(EVENT_TYPE='CLICK') * 100.0
          / NULLIF(COUNT_IF(EVENT_TYPE='SENT'),0), 1)  AS click_rate_pct
FROM QA_FIPSAR_DW.GOLD.FACT_SFMC_ENGAGEMENT;""",
        "insight_callout": (
            "An 88.9% click-to-open rate is industry-leading (benchmark: ~20–25%). This likely reflects "
            "the targeted, consent-based nature of the FIPSAR prospect list — highly qualified audience "
            "relative to typical broadcast campaigns."
        ),
    },

    # ── Q5 ──────────────────────────────────────────────────────────────────
    {
        "num": 5,
        "question": "Which SFMC journey is performing best? Compare open rates and click rates across all journeys.",
        "category": "Journey Performance Comparison",
        "business_insight": (
            "The Prospect Journey has 9 stages grouped into 4 phases. Welcome Phase (Stages 01–02) "
            "typically achieves the highest open rates because prospects are most engaged at initial "
            "outreach. Low Engagement Phase (Stages 08–09: Re-engagement and Final Reminder) has "
            "lower open rates by design — it targets lapsed prospects — but a high click-to-open "
            "ratio when it does land indicates content relevance."
        ),
        "expected_headers": ["Phase", "Stages", "Sent", "Opened", "Clicked", "Open Rate"],
        "expected_rows": [
            ["Welcome Phase",          "01–02", "~1,100", "~870",  "~790",  "~79%"],
            ["Nurture Phase",          "03–05", "~1,200", "~890",  "~800",  "~74%"],
            ["High Engagement Phase",  "06–07", "~986",  "~776",  "~655",  "~79%"],
            ["Low Engagement Phase",   "08–09", "~1,000", "~710",  "~640",  "~71%"],
        ],
        "validation_sql": """\
-- Prospect Journey: phase-level email performance (via DIM_SFMC_JOB JOURNEY_TYPE)
SELECT
    j.JOURNEY_TYPE                                        AS phase_code,
    COUNT_IF(e.EVENT_TYPE = 'SENT')                      AS sent,
    COUNT_IF(e.EVENT_TYPE = 'OPEN')                      AS opened,
    COUNT_IF(e.EVENT_TYPE = 'CLICK')                     AS clicked,
    ROUND(COUNT_IF(e.EVENT_TYPE='OPEN')  * 100.0
          / NULLIF(COUNT_IF(e.EVENT_TYPE='SENT'),0), 1)  AS open_rate_pct,
    ROUND(COUNT_IF(e.EVENT_TYPE='CLICK') * 100.0
          / NULLIF(COUNT_IF(e.EVENT_TYPE='OPEN'),0), 1)  AS click_to_open_pct
FROM QA_FIPSAR_DW.GOLD.FACT_SFMC_ENGAGEMENT e
LEFT JOIN QA_FIPSAR_DW.GOLD.DIM_SFMC_JOB j ON e.JOB_KEY = j.JOB_KEY
GROUP BY j.JOURNEY_TYPE
ORDER BY open_rate_pct DESC;""",
        "insight_callout": (
            "Welcome Phase and High Engagement Phase both hit ~79% open rates. "
            "Prospects who reach Stage 06 — Conversion Email have demonstrated sustained "
            "engagement through 5 prior stages, making them the highest-priority conversion targets."
        ),
    },

    # ── Q6 ──────────────────────────────────────────────────────────────────
    {
        "num": 6,
        "question": "Show me the monthly trend of lead intake for Q1 2026. Is volume growing or declining?",
        "category": "Trend Analysis",
        "business_insight": (
            "Lead intake peaked in February/March 2026, reflecting a seasonal ramp-up typical of "
            "healthcare campaign cycles post-open enrollment. April shows partial data (records only "
            "through Apr 1). Month-over-month growth from January to March indicates the marketing "
            "spend was scaling up effectively through Q1."
        ),
        "expected_headers": ["Month", "Raw Leads", "Valid Prospects", "Conversion %"],
        "expected_rows": [
            ["January 2026",  "~220", "~208", "~94.5%"],
            ["February 2026", "~215", "~204", "~94.9%"],
            ["March 2026",    "~282", "~267", "~94.7%"],
            ["April 2026",    "~83",  "~79",  "~95.2%"],
        ],
        "validation_sql": """\
-- Monthly lead intake trend
SELECT
    TO_CHAR(FILE_DATE, 'YYYY-MM')            AS intake_month,
    COUNT(*)                                  AS raw_leads,
    COUNT_IF(PATIENT_CONSENT = TRUE
             AND EMAIL IS NOT NULL
             AND PHONE_NUMBER IS NOT NULL)    AS clean_leads
FROM QA_FIPSAR_PHI_HUB.STAGING.STG_PROSPECT_INTAKE
GROUP BY intake_month
ORDER BY intake_month;""",
        "insight_callout": (
            "March shows the highest intake volume — typical for Q1 healthcare campaigns closing out "
            "before end-of-quarter. Monitor whether April maintains momentum into Q2."
        ),
    },

    # ── Q7 ──────────────────────────────────────────────────────────────────
    {
        "num": 7,
        "question": "How many prospects have bounced, unsubscribed, or filed spam complaints? What is the suppression rate?",
        "category": "Suppression & Deliverability",
        "business_insight": (
            "Suppression events (bounces, unsubscribes, spam complaints) are critical deliverability "
            "signals. A high hard-bounce rate indicates stale or invalid email addresses in the prospect "
            "pool. Unsubscribes signal journey fatigue — if they spike after JNY002 it means the nurture "
            "sequence is too aggressive. Any spam complaints need immediate review."
        ),
        "expected_headers": ["Event Type", "Event Count", "Unique Prospects", "Rate vs Sent"],
        "expected_rows": [
            ["BOUNCE",         "~180", "~165", "~4.2%"],
            ["UNSUBSCRIBE",    "~95",  "~90",  "~2.2%"],
            ["SPAM_COMPLAINT", "~20",  "~19",  "~0.5%"],
            ["Total Suppressed","~295","~274",  "~6.4%"],
        ],
        "validation_sql": """\
-- Suppression and deliverability event counts
SELECT
    EVENT_TYPE,
    COUNT(*)                           AS event_count,
    COUNT(DISTINCT SUBSCRIBER_KEY)     AS unique_prospects,
    ROUND(COUNT(*) * 100.0 / (
        SELECT COUNT(*) FROM QA_FIPSAR_DW.GOLD.FACT_SFMC_ENGAGEMENT
        WHERE EVENT_TYPE = 'SENT'
    ), 2)                              AS rate_vs_sent_pct
FROM QA_FIPSAR_DW.GOLD.FACT_SFMC_ENGAGEMENT
WHERE EVENT_TYPE IN ('BOUNCE','UNSUBSCRIBE','SPAM_COMPLAINT')
GROUP BY EVENT_TYPE
ORDER BY event_count DESC;""",
        "insight_callout": (
            "If bounce rate exceeds 5%, review email validation at intake. FIPSAR currently rejects "
            "NULL_EMAIL records at PHI stage — but valid-format emails that no longer exist still "
            "pass the gate and will surface as bounces in SFMC."
        ),
    },

    # ── Q8 ──────────────────────────────────────────────────────────────────
    {
        "num": 8,
        "question": "What is the age group distribution of our Gold-layer prospects and which group has the highest email engagement?",
        "category": "Demographic Segmentation",
        "business_insight": (
            "Understanding age demographics helps personalise journey content. Healthcare marketing "
            "typically sees higher engagement from the 56-65 and 65+ cohorts who are actively evaluating "
            "insurance or care options. Younger cohorts (18-35) may require mobile-optimised, shorter "
            "email formats to drive click-through."
        ),
        "expected_headers": ["Age Group", "Prospect Count", "% of Total", "Avg Open Rate"],
        "expected_rows": [
            ["18-25", "~95",  "12.5%", "~72%"],
            ["26-35", "~120", "15.8%", "~74%"],
            ["36-45", "~135", "17.8%", "~76%"],
            ["46-55", "~145", "19.1%", "~78%"],
            ["56-65", "~140", "18.5%", "~80%"],
            ["65+",   "~123", "16.2%", "~81%"],
        ],
        "validation_sql": """\
-- Age group distribution with engagement rates
SELECT
    d.AGE_GROUP,
    COUNT(DISTINCT d.PROSPECT_KEY)                             AS prospect_count,
    ROUND(COUNT(DISTINCT d.PROSPECT_KEY) * 100.0
          / SUM(COUNT(DISTINCT d.PROSPECT_KEY)) OVER(), 1)     AS pct_of_total,
    ROUND(COUNT_IF(e.EVENT_TYPE = 'OPEN') * 100.0
          / NULLIF(COUNT_IF(e.EVENT_TYPE = 'SENT'), 0), 1)     AS open_rate_pct
FROM QA_FIPSAR_DW.GOLD.DIM_PROSPECT          d
LEFT JOIN QA_FIPSAR_DW.GOLD.FACT_SFMC_ENGAGEMENT e
       ON d.MASTER_PATIENT_ID = e.SUBSCRIBER_KEY
GROUP BY d.AGE_GROUP
ORDER BY d.AGE_GROUP;""",
        "insight_callout": (
            "65+ cohort shows the highest open rate (~81%), confirming that older demographics are "
            "more likely to engage with healthcare-related content. Tailor JNY002/JNY004 messaging "
            "to this segment's concerns: coverage, affordability, preventive care."
        ),
    },

    # ── Q9 ──────────────────────────────────────────────────────────────────
    {
        "num": 9,
        "question": "Which states have the highest prospect concentration, and are high-volume states also high-engagement states?",
        "category": "Geographic Analysis",
        "business_insight": (
            "Geographic concentration reveals where marketing spend is working and where there may be "
            "under-served markets. High-volume states that also show high email engagement rates "
            "represent the best ROI markets and should receive increased budget allocation. "
            "States with high volume but low engagement may require localised content or different channels."
        ),
        "expected_headers": ["State", "Prospect Count", "Sent", "Opened", "Open Rate"],
        "expected_rows": [
            ["FL", "~72", "~380", "~295", "~77.6%"],
            ["CA", "~68", "~355", "~270", "~76.1%"],
            ["TX", "~65", "~340", "~258", "~75.9%"],
            ["NY", "~62", "~320", "~245", "~76.6%"],
            ["IL", "~55", "~290", "~222", "~76.6%"],
        ],
        "validation_sql": """\
-- State-level prospect count and engagement
SELECT
    d.STATE,
    COUNT(DISTINCT d.PROSPECT_KEY)                             AS prospect_count,
    COUNT_IF(e.EVENT_TYPE = 'SENT')                           AS sent,
    COUNT_IF(e.EVENT_TYPE = 'OPEN')                           AS opened,
    ROUND(COUNT_IF(e.EVENT_TYPE='OPEN') * 100.0
          / NULLIF(COUNT_IF(e.EVENT_TYPE='SENT'),0), 1)        AS open_rate_pct
FROM QA_FIPSAR_DW.GOLD.DIM_PROSPECT          d
LEFT JOIN QA_FIPSAR_DW.GOLD.FACT_SFMC_ENGAGEMENT e
       ON d.MASTER_PATIENT_ID = e.SUBSCRIBER_KEY
WHERE d.STATE <> 'UNKNOWN'
GROUP BY d.STATE
ORDER BY prospect_count DESC
LIMIT 10;""",
        "insight_callout": (
            "Note: The synthetic dataset has many STATE = 'UNKNOWN' values — in production, "
            "enriching the geo-lookup table (GEO_DIM) with valid USPS state codes will unlock "
            "full regional performance analysis."
        ),
    },

    # ── Q10 ─────────────────────────────────────────────────────────────────
    {
        "num": 10,
        "question": "How many prospects have completed the full 4-journey SFMC sequence vs dropped off — and at which journey stage is churn highest?",
        "category": "Journey Completion & Churn",
        "business_insight": (
            "The Prospect Journey has 9 stages. Stage completion analysis reveals where prospects "
            "drop off. Suppressed prospects are identified via IS_SUPPRESSED=TRUE — NULL stage columns "
            "after the last TRUE stage are intentional suppression cutoffs, not missing data. "
            "The biggest drop typically happens at Stage 03 (Education Email 1) — where prospects "
            "who engaged with Welcome Phase emails but did not re-engage in Nurture Phase fall off."
        ),
        "expected_headers": ["Stage Reached", "Stage Name", "Unique Prospects", "Drop", "Completion %"],
        "expected_rows": [
            ["Stage 01", "Welcome Email",       "758",  "—",    "100%"],
            ["Stage 02", "Education Email",      "~728", "~30",  "~96.0%"],
            ["Stage 03", "Education Email 1",    "~694", "~34",  "~91.6%"],
            ["Stage 04", "Education Email 2",    "~668", "~26",  "~88.1%"],
            ["Stage 05", "Prospect Story Email", "~648", "~20",  "~85.5%"],
            ["Stage 06", "Conversion Email",     "~620", "~28",  "~81.8%"],
            ["Stage 07", "Reminder Email",       "~598", "~22",  "~78.9%"],
            ["Stage 08", "Re-engagement Email",  "~572", "~26",  "~75.5%"],
            ["Stage 09", "Final Reminder Email", "~548", "~24",  "~72.3%"],
        ],
        "validation_sql": """\
-- Prospect Journey stage completion — unique prospects per stage
SELECT
    COUNT_IF(UPPER(TRIM(WELCOMEJOURNEY_WELCOMEEMAIL_SENT))='TRUE')        AS stage_01_welcome,
    COUNT_IF(UPPER(TRIM(WELCOMEJOURNEY_EDUCATIONEMAIL_SENT))='TRUE')      AS stage_02_education,
    COUNT_IF(UPPER(TRIM(NURTUREJOURNEY_EDUCATIONEMAIL1_SENT))='TRUE')     AS stage_03_education_1,
    COUNT_IF(UPPER(TRIM(NURTUREJOURNEY_EDUCATIONEMAIL2_SENT))='TRUE')     AS stage_04_education_2,
    COUNT_IF(UPPER(TRIM(NURTUREJOURNEY_PROSPECTSTORYEMAIL_SENT))='TRUE')  AS stage_05_prospect_story,
    COUNT_IF(UPPER(TRIM(HIGHENGAGEMENT_CONVERSIONEMAIL_SENT))='TRUE')     AS stage_06_conversion,
    COUNT_IF(UPPER(TRIM(HIGHENGAGEMENT_REMINDEREMAIL_SENT))='TRUE')       AS stage_07_reminder,
    COUNT_IF(UPPER(TRIM(LOWENGAGEMENT_REENGAGEMENTEMAIL_SENT))='TRUE')    AS stage_08_reengagement,
    COUNT_IF(UPPER(TRIM(LOWENGAGEMENTFINALREMINDEREMAIL_SENT))='TRUE')    AS stage_09_final_reminder,
    COUNT_IF(UPPER(TRIM(SUPPRESSION_FLAG)) IN ('YES','Y','TRUE','1'))     AS total_suppressed
FROM QA_FIPSAR_SFMC_EVENTS.RAW_EVENTS.RAW_SFMC_PROSPECT_JOURNEY_DETAILS;""",
        "insight_callout": (
            "Prospects who reach Stage 06 — Conversion Email but do not click should be flagged "
            "for manual outreach. They have demonstrated sustained engagement across 5 prior stages — "
            "a personalised touchpoint at this point has the highest likelihood of converting them. "
            "Always compute LAST_COMPLETED_STAGE to understand suppression cutoff point."
        ),
    },
]


# ─────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover / Title block ────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after  = Pt(4)
    add_run(title_para, "FIPSAR Intelligence", bold=True, color=NAVY, size=26, font="Calibri Light")

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(2)
    add_run(sub_para, "Demo Q&A — Business Insights, Expected Outputs & Validation SQL",
            bold=False, color=TEAL, size=13)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_after = Pt(16)
    add_run(meta_para, "QA Environment  |  Q1 2026  |  10 Representative Questions",
            italic=True, color=MID_GRAY, size=10)

    para_border_bottom(meta_para, color="1B3A6B", sz="12")
    doc.add_paragraph()

    # ── Introduction ──────────────────────────────────────────────────────
    heading1(doc, "Overview")
    body_para(doc, (
        "This document presents 10 representative business questions that demonstrate the analytical "
        "depth of the FIPSAR Intelligence platform. Each question includes:"
    ))
    for item in [
        "The business question as a user would ask it in the chatbot",
        "A business insight summary explaining what the data tells us",
        "An expected output table showing approximate results",
        "Validation SQL that can be run directly in Snowflake SnowSight (QA environment)",
        "A key insight callout highlighting the actionable takeaway",
    ]:
        bullet_para(doc, item)

    doc.add_paragraph()
    label_para(doc, "Database Environment:  ", "QA_FIPSAR_PHI_HUB, QA_FIPSAR_DW, QA_FIPSAR_AUDIT, QA_FIPSAR_SFMC_EVENTS")
    label_para(doc, "Data Period:  ",           "January 1, 2026 – April 1, 2026 (Q1 2026)")
    label_para(doc, "Total Raw Leads:  ",        "800 records")
    label_para(doc, "Total Valid Prospects:  ",  "758 Gold-layer records")
    label_para(doc, "SFMC Events:  ",            "~10,627 engagement events (SENT, OPEN, CLICK, BOUNCE, UNSUBSCRIBE, SPAM)")
    doc.add_paragraph()

    # ── Questions ─────────────────────────────────────────────────────────
    for q in QUESTIONS:
        divider(doc)
        doc.add_paragraph()

        # Question number + title
        q_title = doc.add_paragraph()
        q_title.paragraph_format.space_before = Pt(6)
        q_title.paragraph_format.space_after  = Pt(4)
        add_run(q_title, f"Q{q['num']}  ", bold=True, color=GOLD, size=16)
        add_run(q_title, q["question"], bold=True, color=NAVY, size=13)

        # Category badge
        cat_para = doc.add_paragraph()
        cat_para.paragraph_format.space_after = Pt(6)
        add_run(cat_para, f"  Category: ", bold=True, color=MID_GRAY, size=9.5)
        add_run(cat_para, q["category"], bold=False, color=TEAL, size=9.5)

        # Business Insight
        heading2(doc, "Business Insight")
        body_para(doc, q["business_insight"])

        # Expected Output
        heading2(doc, "Expected Output (Approximate)")
        expected_table(doc, q["expected_headers"], q["expected_rows"])

        # Validation SQL
        heading2(doc, "Validation SQL")
        sql_block(doc, q["validation_sql"])

        # Key Insight callout
        insight_box(doc, q["insight_callout"])

        doc.add_paragraph()

    # ── Appendix: Data Model Reference ────────────────────────────────────
    doc.add_page_break()
    heading1(doc, "Appendix — QA Data Model Quick Reference")
    doc.add_paragraph()

    heading2(doc, "Pipeline Flow")
    for step in [
        "STG_PROSPECT_INTAKE (800 raw leads)",
        "  ↓  PHI quality gate: reject NULL_EMAIL (31), NULL_PHONE (4), NO_CONSENT (2)",
        "PHI_PROSPECT_MASTER (763 mastered records)",
        "PATIENT_IDENTITY_XREF (763 identity mappings)",
        "BRZ_PROSPECT_MASTER (763 bronze)",
        "  ↓  Silver deduplication: remove 5 DUPLICATE_RECORD_ID",
        "SLV_PROSPECT_MASTER (758 silver)",
        "DIM_PROSPECT + FACT_PROSPECT_INTAKE (758 gold)",
        "RAW_SFMC_SENT/OPENS/CLICKS/BOUNCES → FACT_SFMC_ENGAGEMENT (~10,627 events)",
        "DQ_REJECTION_LOG (42 total rejections audited)",
    ]:
        bullet_para(doc, step, size=10)

    doc.add_paragraph()
    heading2(doc, "Key Join Keys")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = tbl.rows[0]
    for i, h in enumerate(["Key", "Column Name", "Used In"]):
        set_cell_bg(hrow.cells[i], "1B3A6B")
        p = hrow.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(9.5); run.font.name = "Calibri"

    join_data = [
        ("Prospect Identity", "MASTER_PATIENT_ID = SUBSCRIBER_KEY", "DIM_PROSPECT ↔ FACT_SFMC_ENGAGEMENT"),
        ("Intake Record",     "RECORD_ID",                           "STG ↔ PHI ↔ BRZ ↔ SLV ↔ FACT_INTAKE"),
        ("Surrogate Key",     "PROSPECT_KEY",                        "DIM_PROSPECT ↔ FACT_PROSPECT_INTAKE"),
        ("Journey",           "JOURNEY_TYPE in DIM_SFMC_JOB (J01–J04 = phases of Prospect Journey)", "FACT_SFMC_ENGAGEMENT via JOB_KEY"),
    ]
    for i, (key, col, used) in enumerate(join_data):
        row = tbl.add_row()
        bg = "FFFFFF" if i % 2 == 0 else "EAF2FF"
        for j, val in enumerate([key, col, used]):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, color="CCCCCC", sz="2")
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9); run.font.name = "Calibri"; run.font.color.rgb = DARK_GRAY

    doc.add_paragraph()
    heading2(doc, "Prospect Journey — 9 Stages, 4 Phases")
    for stage_line in [
        "Stage 01 — Welcome Email           (Welcome Phase)",
        "Stage 02 — Education Email          (Welcome Phase)",
        "Stage 03 — Education Email 1        (Nurture Phase)",
        "Stage 04 — Education Email 2        (Nurture Phase)",
        "Stage 05 — Prospect Story Email     (Nurture Phase)",
        "Stage 06 — Conversion Email         (High Engagement Phase)",
        "Stage 07 — Reminder Email           (High Engagement Phase)",
        "Stage 08 — Re-engagement Email      (Low Engagement Phase)",
        "Stage 09 — Final Reminder Email     (Low Engagement Phase)",
    ]:
        bullet_para(doc, stage_line)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_para, "FIPSAR Intelligence Platform  |  QA Environment  |  Confidential",
            italic=True, color=MID_GRAY, size=9)

    # ── Save ──────────────────────────────────────────────────────────────
    out_path = "FIPSAR_Intelligence_Demo_QA.docx"
    doc.save(out_path)
    print(f"Document saved: {out_path}")


if __name__ == "__main__":
    build_document()
